
import frequecny_profile
from docx import Document
from datetime import datetime
from docx.shared import Pt


f = open('../Logs/Logger.txt',"a+")
f.write('The code is run at ReadDoc Script '+datetime.now().strftime('%Y-%m-%d %H:%M:%S')+ '\n')


doc = Document('../freq files/Format File.docx')

table = doc.tables[0]

table.alignment = 1
style = doc.styles['Normal']
font = style.font
font.name = "Times New Roman"
font.size = Pt(9.5)
font.bold = True

list_of_keys = ['date','max_freq','min_freq','avg_freq','below_band','in_the_band','above_band','outside_band','hours_outside_band','fdi']
for row,i in zip(range(3,10), range(7)):
    for col,j in zip(range(10), list_of_keys):
       table.cell(row,col).text = str(frequecny_profile.frequency_profile_data[i][j])
       table.cell(row,col).paragraphs[0].alignment = 1
       

table.cell(3,10).text = str(frequecny_profile.weekly_fdi())
table.cell(3,10).paragraphs[0].alignment = 1      


doc.save('../freq files/Format File1.docx')