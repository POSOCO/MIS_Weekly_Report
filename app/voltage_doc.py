import vol_profile
from docx import Document
from datetime import datetime
from docx.shared import Pt

f = open('../Logs/Logger.txt',"a+")
f.write('The code is run at Vol Doc Script '+datetime.now().strftime('%Y-%m-%d %H:%M:%S')+ '\n')


doc = Document('../freq files/Format File Vol.docx')

table1 = doc.tables[0]
table2 = doc.tables[1]
table3 = doc.tables[2]
table4 = doc.tables[3]

table1.alignment = 1
table2.alignment = 2
table3.alignment = 3
table4.alignment = 4


style = doc.styles['Normal']
font = style.font
font.name = "Times New Roman"
font.size = Pt(10)
font.bold = True


table1Stations = ['Bhopal','Khandwa','Itarsi','Damoh','Nagda','Indore','Gwalior','Raipur','Raigarh']
table2Stations = ['Bhilai','Wardha','Dhule','Parli','Boisar','Kalwa','Karad','Asoj','Dehgam']
table3Stations= ['Kasor','Jetpur','Amreli','Vapi','Sipat','Seoni','Wardh','Bina','Ind765']
table4Stations = ['Sasan','Satna','Tamnar','Kotra','Vododara','Durg','Gwa765']

tablesStations = table1Stations+table2Stations+table3Stations+table4Stations

tables = [table1,table2,table3,table4]
"""i = 0
j = 0
for eachTable in tables:
    i = j
    for row  in range(2,8):
        j = i
        for col in range(1,19):
            if col%2 == 0:
                print('We are in if loop ,col value is '+str(col) +' and j value is '+str(j))
                eachTable.cell(row,col).text = str(vol_profile.vol_profile_data[j][tablesStations[j]]['Min'])
                eachTable.cell(row,col).paragraphs[0].alignment = 1
                #print (tablesStations[j])
                #print(vol_profile.vol_profile_data[j][tablesStations[j]]['Min'])
                j += 1
                if j == len(tablesStations):
                    break
                
                
                #print(vol_profile.vol_profile_data[i][eachStation]['Min'])
            else:
                
                print('We are in else loop ,col value is '+str(col) +' and j value is '+str(j))
                eachTable.cell(row,col).text = str(vol_profile.vol_profile_data[j][tablesStations[j]]['Max'])
                eachTable.cell(row,col).paragraphs[0].alignment = 1
               # print(tablesStations[j])
        
                
               # print(vol_profile.vol_profile_data[j][tablesStations[j]]['Max'])
               # print(vol_profile.vol_profile_data[i][eachStation]['Min'])

           

"""
dateList = vol_profile.dateList


'''for eachStation in table1Staions:
    for row,date in zip(range(2,9),dateList):
        table1.cell(row,0).text = str(date)
        for col in range(1,19):  
            if col%2 == 0:
                print(eachStation)
                table1.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Min'])
                table1.cell(row,col).paragraphs[0].alignment = 1 
            else:
                print(eachStation)
                table1.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Max'])
                table1.cell(row,col).paragraphs[0].alignment = 1 
'''
for row,date in zip(range(2,9),dateList):
        table1.cell(row,0).text = str(date) 
        table2.cell(row,0).text = str(date) 
        table3.cell(row,0).text = str(date) 
        table4.cell(row,0).text = str(date) 
        station = 0
        for col in range(1,19):
            if col%2 == 0:
                
                table1.cell(row,col).text = str(vol_profile.vol_profile_data[table1Stations[station]][date]['Min'])
                table2.cell(row,col).text = str(vol_profile.vol_profile_data[table2Stations[station]][date]['Min'])
                table3.cell(row,col).text = str(vol_profile.vol_profile_data[table3Stations[station]][date]['Min'])
                table1.cell(row,col).paragraphs[0].alignment = 1 
                table2.cell(row,col).paragraphs[0].alignment = 1 
                table3.cell(row,col).paragraphs[0].alignment = 1 

                station += 1
            else:
                table1.cell(row,col).text = str(vol_profile.vol_profile_data[table1Stations[station]][date]['Max'])
                table2.cell(row,col).text = str(vol_profile.vol_profile_data[table2Stations[station]][date]['Max'])
                table3.cell(row,col).text = str(vol_profile.vol_profile_data[table3Stations[station]][date]['Max'])
                table1.cell(row,col).paragraphs[0].alignment = 1 
                table2.cell(row,col).paragraphs[0].alignment = 1 
                table3.cell(row,col).paragraphs[0].alignment = 1
            
'''for eachStation in table2Stations:
    for row,date in zip(range(2,9),dateList):
        table2.cell(row,0).text = str(date)
        for col in range(1,19):  
            if col%2 == 0:
                table2.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Min'])
                table2.cell(row,col).paragraphs[0].alignment = 1 
            else:
                table2.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Max'])
                table2.cell(row,col).paragraphs[0].alignment = 1 

for eachStation in table3Stations:
    for row,date in zip(range(2,9),dateList):
        table3.cell(row,0).text = str(date)
        for col in range(1,19):  
            if col%2 == 0:
                table3.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Min'])
                table3.cell(row,col).paragraphs[0].alignment = 1 
            else:
                table3.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Max'])
                table3.cell(row,col).paragraphs[0].alignment = 1 

for eachStation in table4Stations:
    for row,date in zip(range(2,9),dateList):
        table4.cell(row,0).text = str(date)
        for col in range(1,14):  
            if col%2 == 0:
                table4.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Min'])
                table4.cell(row,col).paragraphs[0].alignment = 1 
            else:
                table4.cell(row,col).text = str(vol_profile.vol_profile_data[eachStation][date]['Max'])
                table4.cell(row,col).paragraphs[0].alignment = 1 

'''
doc.save('../freq files/Format File Vol1.docx')
